#!/usr/bin/env python3
"""Genera documentación académica en Word (.docx) desde Markdown."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "docs" / "DOCUMENTACION_DEL_SISTEMA.md"
OUTPUT = ROOT / "docs" / "Documentacion_del_Sistema_PawPatrol.docx"

# Paleta institucional
AZUL = RGBColor(0x1A, 0x3A, 0x6E)
AZUL_CLARO = RGBColor(0x2E, 0x6B, 0xA8)
GRIS = RGBColor(0x55, 0x55, 0x55)
GRIS_FONDO = "F2F5FA"
GRIS_CODIGO = "F4F4F4"
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

DIAGRAMAS_LEGIBLES: dict[str, str] = {
    "contexto": (
        "DIAGRAMA DE CONTEXTO — NIVEL 1\n\n"
        "Actores:\n"
        "  • Usuario registrado — fichas, pérdidas, avistamientos, chat y búsqueda por foto.\n"
        "  • Administrador (único) — panel /admin, estadísticas y moderación.\n\n"
        "Sistema central:\n"
        "  • Aplicación web PawPatrol (Next.js 16).\n"
        "  • Servicio WebSocket (Railway) para tiempo real.\n\n"
        "Servicios externos:\n"
        "  • Neon PostgreSQL — base de datos.\n"
        "  • Vercel Blob — fotos y adjuntos.\n"
        "  • Gmail SMTP — correos transaccionales.\n"
        "  • Google — OAuth, Maps/Places, Gemini IA."
    ),
    "secuencia": (
        "DIAGRAMA DE SECUENCIA — PÉRDIDA Y AVISTAMIENTO\n\n"
        "Nota: Dueño y Testigo son papeles del CASO (mascota.user_id / avistamiento.user_id),\n"
        "no roles de cuenta. Ambos suelen ser USUARIO en la BD.\n\n"
        "1. Usuario A (dueño del caso) marca mascota PERDIDA.\n"
        "2. Sistema actualiza BD e indexa embedding visual (Gemini).\n"
        "3. Usuario B (testigo) reporta avistamiento vinculado.\n"
        "4. Sistema guarda avistamiento PENDIENTE y notifica a A.\n"
        "5. Evento WebSocket actualiza mapa y chat.\n"
        "6. A verifica el avistamiento; chat privado A ↔ B.\n"
        "7. Mensajes y lecturas quedan en PostgreSQL."
    ),
    "mer": (
        "MODELO ENTIDAD-RELACIÓN (RESUMEN)\n\n"
        "USER (1) —— (N) MASCOTA\n"
        "USER (1) —— (N) AVISTAMIENTO\n"
        "USER (1) —— (N) NOTIFICACION\n"
        "MASCOTA (1) —— (N) MASCOTA_FOTO\n"
        "MASCOTA (1) —— (N) MASCOTA_EMBEDDING\n"
        "MASCOTA (1) —— (N) AVISTAMIENTO\n"
        "MASCOTA (1) —— (N) HISTORIAL_ESTADO\n"
        "MASCOTA (1) —— (N) EVENTO_CASO\n"
        "AVISTAMIENTO (1) —— (N) MENSAJE_AVISTAMIENTO\n"
        "AVISTAMIENTO (1) —— (N) LECTURA_CHAT\n"
        "AVISTAMIENTO (1) —— (N) REPORTE_ABUSO"
    ),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)


def add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()

    run = p.add_run("PawPatrol — Documentación del Sistema  |  Página ")
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_run = OxmlElement("w:r")
    fld_run.append(fld_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = p.add_run()
    r._r.append(fld_begin)
    r2 = p.add_run()
    r2._r.append(instr)
    r3 = p.add_run()
    r3._r.append(fld_sep)
    r3._r.append(fld_run)
    r4 = p.add_run()
    r4._r.append(fld_end)


def add_header(section) -> None:
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.clear()
    run = p.add_run("PAWPATROLL — Documentación del Sistema v1.0")
    run.font.size = Pt(9)
    run.font.color.rgb = AZUL_CLARO
    run.italic = True


def add_toc(doc: Document) -> None:
    doc.add_heading("Índice general", level=1)
    p = doc.add_paragraph()
    run = p.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run2 = p.add_run()
    run2._r.append(instr)
    run3 = p.add_run()
    run3._r.append(fld_char_sep)
    t = OxmlElement("w:t")
    t.text = "Actualice el índice en Word: clic derecho → Actualizar campo."
    run3._r.append(t)
    run4 = p.add_run()
    run4._r.append(fld_char_end)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    r = note.add_run(
        "Nota: al abrir en Microsoft Word, pulse clic derecho sobre el índice "
        "y seleccione «Actualizar campo» → «Actualizar toda la tabla»."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    band = doc.add_paragraph()
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(band, "1A3A6E")
    band.paragraph_format.space_before = Pt(0)
    band.paragraph_format.space_after = Pt(0)
    r = band.add_run("   ")
    r.font.size = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("DOCUMENTACIÓN DEL SISTEMA")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = AZUL

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    r = subtitle.add_run("PawPatrol (PAWPATROLL)")
    r.font.size = Pt(18)
    r.font.color.rgb = AZUL_CLARO

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_after = Pt(24)
    r = desc.add_run(
        "Plataforma comunitaria para reportar pérdidas, registrar avistamientos\n"
        "y gestionar fichas digitales de mascotas con mapa, IA y chat en tiempo real"
    )
    r.font.size = Pt(12)
    r.font.color.rgb = GRIS

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run("─" * 42)
    r.font.color.rgb = AZUL_CLARO

    meta_items = [
        ("Versión", "1.0"),
        ("Fecha", "Junio 2026"),
        ("Autor", "Branly Paucar Arias"),
        ("Usuario GitHub", "Smith2207"),
        ("Demo", "pawpatroll.vercel.app"),
        ("Repositorio", "github.com/Smith2207/PAWPATROLL"),
    ]
    table = doc.add_table(rows=len(meta_items), cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (k, v) in enumerate(meta_items):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.width = Inches(1.8)
        c1.width = Inches(3.5)
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rk = p0.add_run(k + ":")
        rk.bold = True
        rk.font.size = Pt(11)
        rk.font.color.rgb = AZUL
        p1 = c1.paragraphs[0]
        rv = p1.add_run(v)
        rv.font.size = Pt(11)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_before = Pt(36)
    r = foot.add_run(
        "Documento elaborado según estructura académica\n"
        "«Documentación del Sistema — Versión Concisa y Resumida»"
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRIS

    doc.add_page_break()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for level, size, color in [
        ("Heading 1", 16, AZUL),
        ("Heading 2", 13, AZUL_CLARO),
        ("Heading 3", 12, RGBColor(0x33, 0x33, 0x33)),
    ]:
        style = doc.styles[level]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E6BA8")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run_el.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run_el.append(text_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, text: str) -> None:
    text = text.replace("→", "→").replace("↔", "↔")
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|✅|⚠️|☐)"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        chunk = match.group(0)
        if chunk == "✅":
            run = paragraph.add_run("✓ ")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x7A, 0x3C)
        elif chunk == "⚠️":
            run = paragraph.add_run("⚠ ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xC6, 0x7A, 0x00)
        elif chunk == "☐":
            paragraph.add_run("☐ ")
        elif chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        elif chunk.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            if m:
                url = m.group(2)
                if url.startswith("../"):
                    url = url.replace("../", "")
                if not url.startswith("http"):
                    url = f"https://github.com/Smith2207/PAWPATROLL/blob/main/{url}"
                add_hyperlink(paragraph, m.group(1), url)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells if c)


def normalize_rows(rows: list[list[str]]) -> list[list[str]]:
    cols = max(len(r) for r in rows)
    return [r + [""] * (cols - len(r)) for r in rows]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    rows = normalize_rows(rows)
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLANCO
        set_cell_shading(cell, "1A3A6E")

    for i, row in enumerate(body, start=1):
        fill = "FFFFFF" if i % 2 else GRIS_FONDO
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            add_rich_text(cell.paragraphs[0], cell_text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            set_cell_shading(cell, fill)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def detect_diagram(code_lang: str, code_lines: list[str]) -> str | None:
    blob = "\n".join(code_lines).lower()
    if code_lang != "mermaid":
        return None
    if "flowchart tb" in blob and "neon" in blob:
        return DIAGRAMAS_LEGIBLES["contexto"]
    if "sequencediagram" in blob:
        return DIAGRAMAS_LEGIBLES["secuencia"]
    if "erdiagram" in blob:
        return DIAGRAMAS_LEGIBLES["mer"]
    return (
        "DIAGRAMA ARQUITECTÓNICO\n\n"
        + "\n".join(code_lines)
        .replace("flowchart TB", "")
        .replace("sequenceDiagram", "")
        .strip()
    )


def add_diagram_box(doc: Document, content: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_shading(p, GRIS_FONDO)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "1A3A6E")
    p_bdr.append(left)
    p_pr.append(p_bdr)
    run = p.add_run(content)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, GRIS_CODIGO)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def should_skip_line(line: str, skip_meta: bool) -> bool:
    if not skip_meta:
        return False
    if line.startswith("**Versión:**") or line.startswith("**Autor:**"):
        return True
    if line.startswith("**Demo:**") or line.startswith("**Repositorio:**"):
        return True
    if line.startswith("**Fecha:**"):
        return True
    if line.startswith("> Documento elaborado"):
        return True
    if line.strip() == "---":
        return True
    if line.startswith("## Índice"):
        return True
    if re.match(r"^\d+\. \[", line):
        return True
    return False


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    configure_styles(doc)
    add_header(doc.sections[0])
    add_page_number_footer(doc.sections[0])
    add_cover(doc)
    add_toc(doc)

    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_table = False
    skip_until_section = True
    first_h2 = True

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                in_code = False
                diagram = detect_diagram(code_lang, code_lines)
                if diagram:
                    add_diagram_box(doc, diagram)
                else:
                    add_code_block(doc, code_lines)
                code_lang = ""
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            cells = parse_table_row(line)
            if is_separator_row(cells):
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        elif in_table:
            add_table(doc, table_rows)
            in_table = False
            table_rows = []

        if should_skip_line(line, skip_until_section):
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            continue  # título ya en portada

        if line.startswith("## "):
            title = line[3:].strip()
            if title.lower() == "índice":
                continue
            skip_until_section = False
            if not first_h2:
                doc.add_page_break()
            first_h2 = False
            doc.add_heading(title, level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(p, "EEF3FA")
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = GRIS
            continue

        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:].strip())
            continue

        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\.\s*", "", line))
            continue

        if line.strip() == "---":
            continue

        if line.startswith("*Documento generado"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            run = p.add_run(line.strip("*"))
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = GRIS
            continue

        p = doc.add_paragraph()
        add_rich_text(p, line)

    if in_table and table_rows:
        add_table(doc, table_rows)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"✓ Documento Word generado: {out_path}")
    print("  Abra en Microsoft Word y actualice el índice (clic derecho → Actualizar campo).")


if __name__ == "__main__":
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else INPUT
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT
    convert(src, dst)
